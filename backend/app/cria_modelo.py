from docx import Document

# Cria um documento em branco
doc = Document()

# Adiciona o texto com as tags que o nosso sistema Alinhave vai ler
doc.add_heading('DECLARAÇÃO DE CADASTRO - ALINHAVE', level=1)

doc.add_paragraph(
    "Eu, {{ nome }}, portador(a) do CPF nº {{ cpf }}, residente e domiciliado(a) no endereço {{ endereco }}, "
    "com o telefone de contato {{ telefone }} e e-mail {{ email }}, declaro para os devidos fins que as "
    "informações cadastradas no sistema Alinhave são verdadeiras."
)

doc.add_paragraph("\nGoiânia, 07 de Março de 2026.\n")

doc.add_paragraph("_________________________________________________")
doc.add_paragraph("{{ nome }}")

# Salva o arquivo na sua pasta
doc.save('modelo_teste.docx')

print("✅ Arquivo 'modelo_teste.docx' criado com sucesso na sua pasta!")